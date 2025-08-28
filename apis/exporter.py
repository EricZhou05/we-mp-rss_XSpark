import io
import datetime
import pytz
import zipfile  # 引入zipfile库
import json
from typing import Optional, List
from urllib.parse import quote

import docx
from docx.shared import Inches, Pt
from docx.enum.text import WD_LINE_SPACING
from docx.oxml.ns import qn
from docx.oxml import OxmlElement

from fastapi import APIRouter, Depends, Query, HTTPException, status
from fastapi.responses import StreamingResponse
from sqlalchemy.orm import Session

from core.auth import get_current_user
from core.db import DB
from core.models import Feed, Article
from core.models.tags import Tags
from apis.base import error_response

# --- 全局配置 ---
router = APIRouter(prefix="/exporter", tags=["导出选题"])
beijing_tz = pytz.timezone('Asia/Shanghai')
utc_tz = pytz.utc


# --- 辅助函数 ---

# --- 新增：文件名清理与长度限制 ---
INVALID_FILENAME_CHARS = set('\\/:*?"<>|')
MAX_FILENAME_LEN = 150  # 你可以按需调整

def sanitize_filename(component: str) -> str:
    """替换非法字符，去除尾部空格/点，避免跨平台问题。"""
    s = ''.join('_' if c in INVALID_FILENAME_CHARS else c for c in component)
    s = s.strip().rstrip('.')  # Windows不允许以点结束
    return s or '未命名'

def build_safe_filename(export_name: str, first_time: datetime.datetime, last_time: datetime.datetime, max_len: int = MAX_FILENAME_LEN) -> str:
    """拼接并限制文件名长度，过长则截断并加省略号。"""
    prefix = "星火选题库_"
    date_part = f"({first_time.strftime('%m.%d')}_{last_time.strftime('%m.%d')})"
    ext = ".docx"

    export_name = sanitize_filename(export_name)
    filename = f"{prefix}{export_name}{date_part}{ext}"

    if len(filename) <= max_len:
        return filename

    remain = max_len - len(prefix) - len(date_part) - len(ext)
    if remain <= 1:
        # 极端情况：仅保留“星火选题库(时间段).docx”
        return f"星火选题库{date_part}{ext}"

    ellipsis = "…"
    truncated = export_name[:max(1, remain - len(ellipsis))] + ellipsis
    return f"{prefix}{truncated}{date_part}{ext}"

# --- 新增函数：从DOCX文件中移除缩略图 ---
def remove_thumbnail_from_docx(docx_stream):
    """
    将DOCX文件流作为ZIP包处理，移除其中的缩略图文件。

    Args:
        docx_stream (io.BytesIO): 包含原始DOCX文件内容的内存流。

    Returns:
        io.BytesIO: 不含缩略图的新DOCX文件内容的内存流。
    """
    # 确保流的指针在开头
    docx_stream.seek(0)

    # 创建一个新的内存流用于存放修改后的文件
    final_stream = io.BytesIO()

    # 以只读模式打开原始流，以写入模式打开新流
    with zipfile.ZipFile(docx_stream, 'r') as zin:
        with zipfile.ZipFile(final_stream, 'w', zipfile.ZIP_DEFLATED) as zout:
            # 遍历原始ZIP包中的所有文件
            for item in zin.infolist():
                # 如果文件名不是缩略图，则将其复制到新的ZIP包中
                # 缩略图通常位于 'docProps/thumbnail.*'
                if not item.filename.startswith('docProps/thumbnail'):
                    zout.writestr(item, zin.read(item.filename))

    final_stream.seek(0)
    return final_stream


def set_modern_compatibility(document):
    """修改文档兼容性设置，使其符合最新的Word标准。"""
    settings = document.settings.element
    compat_elements = settings.xpath('w:compat')
    compat = compat_elements[0] if compat_elements else OxmlElement('w:compat')

    compat_setting = OxmlElement('w:compatSetting')
    compat_setting.set(qn('w:name'), 'compatibilityMode')
    compat_setting.set(qn('w:uri'), 'http://schemas.microsoft.com/office/word')
    compat_setting.set(qn('w:val'), '16')
    compat.append(compat_setting)

    if not compat_elements:
        settings.append(compat)


def add_hyperlink(paragraph, text, url):
    """在段落中添加一个格式化的超链接。"""
    part = paragraph.part
    r_id = part.relate_to(url, docx.opc.constants.RELATIONSHIP_TYPE.HYPERLINK, is_external=True)
    hyperlink = OxmlElement('w:hyperlink')
    hyperlink.set(qn('r:id'), r_id)
    new_run = OxmlElement('w:r')
    rPr = OxmlElement('w:rPr')
    font_name = '新宋体'
    rFonts = OxmlElement('w:rFonts')
    rFonts.set(qn('w:ascii'), font_name);
    rFonts.set(qn('w:hAnsi'), font_name);
    rFonts.set(qn('w:eastAsia'), font_name)
    rPr.append(rFonts)
    color = OxmlElement('w:color');
    color.set(qn('w:val'), '0000FF');
    rPr.append(color)
    underline = OxmlElement('w:u');
    underline.set(qn('w:val'), 'single');
    rPr.append(underline)
    new_run.append(rPr)
    new_run.text = text
    hyperlink.append(new_run)
    paragraph._p.append(hyperlink)
    return hyperlink


def create_modern_docx(articles_data):
    """一个工厂函数，用于创建一个配置完整的、现代化的DOCX文档对象。"""
    document = docx.Document()
    set_modern_compatibility(document)
    now_utc = datetime.datetime.now(utc_tz)
    core_properties = document.core_properties
    core_properties.author = '星火调研易';
    core_properties.last_modified_by = '星火调研易'
    core_properties.created = now_utc;
    core_properties.modified = now_utc
    core_properties.comments = '';
    core_properties.title = f"星火选题库导出 {datetime.datetime.now(beijing_tz).strftime('%Y-%m-%d')}"
    style = document.styles['Normal']
    font = style.font;
    font.name = '新宋体';
    style.element.rPr.rFonts.set(qn('w:eastAsia'), '新宋体')
    p_format = style.paragraph_format
    p_format.first_line_indent = Pt(0);
    p_format.left_indent = Pt(0)
    p_format.space_after = Pt(6);
    p_format.line_spacing_rule = WD_LINE_SPACING.SINGLE
    section = document.sections[0]
    section.top_margin = Inches(0.5);
    section.bottom_margin = Inches(0.5)
    section.left_margin = Inches(0.5);
    section.right_margin = Inches(0.5)
    for article, feed in articles_data:
        p_title = document.add_paragraph();
        p_title.add_run(f"{article.title} ({feed.mp_name})").bold = True
        publish_time = datetime.datetime.fromtimestamp(article.publish_time, tz=beijing_tz)
        document.add_paragraph(publish_time.strftime('%Y-%m-%d %H:%M:%S'))
        p_link = document.add_paragraph();
        add_hyperlink(p_link, article.url, article.url)
    return document


# --- FastAPI 路由 ---
@router.get("/docx", summary="按时间范围导出公众号文章为DOCX（支持按公众号或多个标签）")
async def export_articles_to_docx(
        start_date: str = Query(..., description="开始日期 (北京时间, 格式: YYYY-MM-DD HH:MM:SS)"),
        end_date: str = Query(..., description="结束日期 (北京时间, 格式: YYYY-MM-DD HH:MM:SS)"),
        feed_id: Optional[str] = Query(None, description="要导出的单个公众号ID (与 tag_id 二选一), 'all' 表示全部"),
        # 核心：支持同名参数多次出现，例如 ?tag_id=1&tag_id=2
        tag_id: Optional[List[str]] = Query(
            None,
            description="要导出的标签ID，可重复传参: ?tag_id=1&tag_id=2"
        ),
        current_user: dict = Depends(get_current_user),
        db: Session = Depends(DB.session_dependency)
):
    # 校验：二选一（feed_id 或 tag_id），且不能为空
    if (feed_id is None and not tag_id) or (feed_id is not None and tag_id):
        raise HTTPException(status_code=400, detail="必须提供 feed_id 或 tag_id，且二者只能提供一个。")

    try:
        start_dt_naive = datetime.datetime.strptime(start_date, '%Y-%m-%d %H:%M:%S')
        end_dt_naive = datetime.datetime.strptime(end_date, '%Y-%m-%d %H:%M:%S')
        start_dt_beijing = beijing_tz.localize(start_dt_naive)
        end_dt_beijing = beijing_tz.localize(end_dt_naive)
        start_timestamp = int(start_dt_beijing.timestamp())
        end_timestamp = int(end_dt_beijing.timestamp())
    except ValueError:
        raise HTTPException(status_code=400, detail="日期格式错误，应为 'YYYY-MM-DD HH:MM:SS'")

    query = db.query(Article, Feed).join(Feed, Article.mp_id == Feed.id)
    export_name = ""

    # 多个标签
    if tag_id:
        # 查出所有标签，并校验是否有不存在的ID
        tags = db.query(Tags).filter(Tags.id.in_(tag_id)).all()
        found_ids = {str(t.id) for t in tags}
        missing = [tid for tid in tag_id if str(tid) not in found_ids]
        if missing:
            raise HTTPException(
                status_code=status.HTTP_404_NOT_FOUND,
                detail={"code": 40402, "message": f"以下标签ID不存在: {', '.join(map(str, missing))}"}
            )

        # 合并多个标签的公众号ID（去重）
        mps_ids_set = set()
        for t in tags:
            if t.mps_id:
                try:
                    data = json.loads(t.mps_id)
                    # 兼容 [{id: xxx}, ...] 或 [id1, id2]
                    if isinstance(data, list):
                        for item in data:
                            if isinstance(item, dict) and 'id' in item:
                                mps_ids_set.add(str(item['id']))
                            else:
                                mps_ids_set.add(str(item))
                except Exception:
                    # 如果某个标签的 mps_id 不是合法JSON，直接忽略该标签的mps
                    pass

        if not mps_ids_set:
            raise HTTPException(
                status_code=status.HTTP_404_NOT_FOUND,
                detail=f"所选标签下未关联任何公众号"
            )

        query = query.filter(Feed.id.in_(mps_ids_set))
        # 文件名中显示多个标签名，使用 '_' 连接
        export_name = '_'.join([t.name for t in tags])

    elif feed_id is not None:
        if feed_id.lower() == "all":
            export_name = "全部选题"
        else:
            feed = db.query(Feed).filter(Feed.id == feed_id).first()
            if not feed:
                raise HTTPException(
                    status_code=status.HTTP_404_NOT_FOUND,
                    detail={"code": 40401, "message": f"公众号 (ID: {feed_id}) 不存在"}
                )
            query = query.filter(Article.mp_id == feed_id)
            export_name = feed.mp_name

    articles_with_feed = query.filter(Article.publish_time.between(start_timestamp, end_timestamp)) \
        .order_by(Article.publish_time.asc()).all()

    if not articles_with_feed:
        raise HTTPException(status_code=status.HTTP_404_NOT_FOUND, detail="在指定的时间范围和条件下未找到任何文章")

    document = create_modern_docx(articles_with_feed)
    initial_stream = io.BytesIO()
    document.save(initial_stream)
    final_stream = remove_thumbnail_from_docx(initial_stream)

    first_time = datetime.datetime.fromtimestamp(articles_with_feed[0][0].publish_time, tz=beijing_tz)
    last_time = datetime.datetime.fromtimestamp(articles_with_feed[-1][0].publish_time, tz=beijing_tz)
    filename = build_safe_filename(export_name, first_time, last_time)

    headers = {
        'Content-Disposition': f"attachment; filename*=UTF-8''{quote(filename)}",
        'Content-Type': 'application/vnd.openxmlformats-officedocument.wordprocessingml.document'
    }
    return StreamingResponse(final_stream, headers=headers)